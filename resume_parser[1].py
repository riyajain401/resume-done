"""
resume_parser.py
-----------------
Extracts plain text from uploaded resumes (PDF / DOCX / TXT) and from
pasted job descriptions. Also does light structural cleanup so the
RAG chunker gets reasonably clean text to work with.
"""

from __future__ import annotations
import io
import re
from pypdf import PdfReader
from docx import Document


def extract_text_from_pdf(file_bytes: bytes) -> str:
    reader = PdfReader(io.BytesIO(file_bytes))
    pages = []
    for page in reader.pages:
        pages.append(page.extract_text() or "")
    return "\n".join(pages)


def extract_text_from_docx(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def extract_text(filename: str, file_bytes: bytes) -> str:
    """Dispatch based on file extension."""
    lower = filename.lower()
    if lower.endswith(".pdf"):
        raw = extract_text_from_pdf(file_bytes)
    elif lower.endswith(".docx"):
        raw = extract_text_from_docx(file_bytes)
    elif lower.endswith(".txt"):
        raw = file_bytes.decode("utf-8", errors="ignore")
    else:
        raise ValueError(f"Unsupported file type: {filename}")
    return clean_text(raw)


def clean_text(text: str) -> str:
    """Collapse whitespace, strip control characters, keep it readable."""
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def chunk_text(text: str, chunk_size: int = 700, overlap: int = 120) -> list[str]:
    """
    Simple sliding-window chunker over characters. Good enough for
    resume/job-description length documents and keeps the RAG layer
    dependency-free (no need for a heavier tokenizer-aware splitter).
    """
    if len(text) <= chunk_size:
        return [text] if text.strip() else []

    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunk = text[start:end]
        chunks.append(chunk)
        start = end - overlap
    return [c.strip() for c in chunks if c.strip()]
